import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def criar_template_limpo():
    # 1. Garante que a pasta assets existe
    if not os.path.exists("assets"):
        os.makedirs("assets")

    doc = Document()
    
    # Estilos básicos
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- TÍTULO ---
    h = doc.add_heading('CONTRATO DE PRESTAÇÃO DE SERVIÇOS', 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    # --- DADOS DO ALUNO ---
    p = doc.add_paragraph()
    p.add_run('CONTRATANTE: ').bold = True
    p.add_run('{{ nome_aluno }}')
    
    p = doc.add_paragraph()
    p.add_run('CPF: ').bold = True
    p.add_run('{{ cpf_aluno }}  ')
    p.add_run('RG: ').bold = True
    p.add_run('{{ rg_aluno }}')
    
    p = doc.add_paragraph()
    p.add_run('ENDEREÇO: ').bold = True
    p.add_run('{{ endereco_aluno }} - {{ cidade_aluno }} (CEP: {{ cep_aluno }})')

    doc.add_paragraph('')

    # --- DADOS DO CURSO ---
    p = doc.add_paragraph()
    p.add_run('OBJETO: ').bold = True
    p.add_run('Curso de Pós-Graduação em {{ nome_curso }} (Turma: {{ turma }}).')
    p.add_run(' Carga Horária: {{ carga_horaria }}h.')
    
    p = doc.add_paragraph()
    p.add_run('Período: ').bold = True
    p.add_run('{{ data_inicio }} a {{ data_fim }}.')

    doc.add_paragraph('')

    # --- FINANCEIRO (RESUMO) ---
    doc.add_heading('Condições Financeiras', level=1)
    
    table_fin = doc.add_table(rows=1, cols=3)
    table_fin.style = 'Table Grid'
    row = table_fin.rows[0]
    row.cells[0].text = 'Valor Bruto: {{ valor_bruto }}'
    row.cells[1].text = 'Desc: {{ desconto_perc }}'
    row.cells[2].text = 'Total: {{ valor_final }}'

    doc.add_paragraph('')

    # --- TABELA 1: ENTRADA (AQUI ESTAVA O ERRO) ---
    doc.add_heading('1. Detalhamento da Entrada', level=2)
    p = doc.add_paragraph('Pagamento da entrada de {{ entrada_total }} realizado da seguinte forma:')

    # Cria tabela 2x4
    table = doc.add_table(rows=2, cols=4)
    table.style = 'Table Grid'
    
    # Cabeçalho
    hdr = table.rows[0].cells
    hdr[0].text = 'Parc.'
    hdr[1].text = 'Vencimento'
    hdr[2].text = 'Valor'
    hdr[3].text = 'Forma'
    
    # Linha do Loop (Sintaxe CORRETA: tr for ... tr endfor)
    row = table.rows[1].cells
    
    # Célula 1: Abre o loop
    p = row[0].paragraphs[0]
    p.add_run('{% tr for item in tbl_entrada %}{{ item.numero }}')
    
    # Células do meio
    row[1].text = '{{ item.data_vencimento }}'
    row[2].text = '{{ item.valor }}'
    
    # Célula 4: Fecha o loop com TR ENDFOR
    p = row[3].paragraphs[0]
    p.add_run('{{ item.forma_pagamento }}{% tr endfor %}')

    doc.add_paragraph('')

    # --- TABELA 2: SALDO ---
    doc.add_heading('2. Detalhamento do Saldo', level=2)
    p = doc.add_paragraph('O saldo restante de {{ saldo_total }} será parcelado em {{ saldo_qtd }} vezes:')

    table2 = doc.add_table(rows=2, cols=4)
    table2.style = 'Table Grid'
    
    hdr2 = table2.rows[0].cells
    hdr2[0].text = 'Parc.'
    hdr2[1].text = 'Vencimento'
    hdr2[2].text = 'Valor'
    hdr2[3].text = 'Forma'
    
    row2 = table2.rows[1].cells
    
    # Loop Saldo
    p = row2[0].paragraphs[0]
    p.add_run('{% tr for item in tbl_saldo %}{{ item.numero }}')
    
    row2[1].text = '{{ item.data_vencimento }}'
    row2[2].text = '{{ item.valor }}'
    
    p = row2[3].paragraphs[0]
    p.add_run('{{ item.forma_pagamento }}{% tr endfor %}')

    # --- ASSINATURAS ---
    doc.add_paragraph('')
    doc.add_paragraph('')
    doc.add_paragraph('Porto Alegre, {{ data_hoje }}').alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph('')
    doc.add_paragraph('')
    
    p = doc.add_paragraph('___________________________________________')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('{{ nome_aluno }}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph('CPF: {{ cpf_aluno }}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Salva
    path = "assets/modelo_contrato_V2.docx"
    doc.save(path)
    print(f"✅ Template NOVO criado com sucesso em: {path}")

if __name__ == "__main__":
    criar_template_limpo()
