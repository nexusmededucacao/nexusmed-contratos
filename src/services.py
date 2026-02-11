import streamlit as st
import os
import io
import smtplib
import subprocess
import pytz
from datetime import datetime
from docxtpl import DocxTemplate
from docx import Document 
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from src.db import supabase

TEMPLATE_PATH = "assets/modelo_contrato_V2.docx"

# --- HELPERS ---
def format_moeda(valor):
    if valor is None: return "0,00" # Retorna sem R$ para o template controlar
    try:
        if isinstance(valor, str):
            valor = float(valor.replace("R$", "").replace(".", "").replace(",", ".").strip())
        return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
    except: return "0,00"

def format_data(data_obj):
    try:
        if isinstance(data_obj, str):
            data_obj = datetime.strptime(data_obj, "%Y-%m-%d")
        return data_obj.strftime("%d/%m/%Y")
    except: return str(data_obj)

def parse_float(valor):
    try:
        if isinstance(valor, (float, int)): return float(valor)
        return float(str(valor).replace("R$", "").replace(".", "").replace(",", ".").strip())
    except: return 0.0

# --- FUNÇÃO DE PREENCHIMENTO (APENAS PAGAMENTOS) ---
def preencher_tabelas_pagamento(docx_path, dados_entrada, dados_saldo):
    """
    Preenche APENAS as tabelas de Entrada e Saldo via código.
    A tabela de Produto (Índice 1) será ignorada aqui pois já foi preenchida pelo Jinja (Word).
    """
    doc = Document(docx_path)
    
    def aplicar_estilo(run):
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0, 0, 0)

    # 1. PREENCHER ENTRADA (Tabela Índice 2 - AJUSTE SE NECESSÁRIO)
    try:
        tbl_ent = doc.tables[2]
        tbl_ent.style = 'Table Grid'
        
        # Limpa linhas antigas (mantém cabeçalho)
        while len(tbl_ent.rows) > 1:
            tbl_ent._element.remove(tbl_ent.rows[-1]._element)
            
        for item in dados_entrada:
            row = tbl_ent.add_row()
            vals = [str(item['numero']), format_data(item['data']), "R$ " + format_moeda(item['valor']), str(item['forma'])]
            for i, v in enumerate(vals):
                row.cells[i].text = v
                for p in row.cells[i].paragraphs: 
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if p.runs: aplicar_estilo(p.runs[0])
                    else: 
                        r = p.add_run(v)
                        aplicar_estilo(r)
    except Exception as e:
        print(f"Erro Tabela Entrada: {e}")

    # 2. PREENCHER SALDO (Tabela Índice 3 - AJUSTE SE NECESSÁRIO)
    try:
        tbl_sal = doc.tables[3]
        tbl_sal.style = 'Table Grid'
        
        while len(tbl_sal.rows) > 1:
            tbl_sal._element.remove(tbl_sal.rows[-1]._element)
            
        for item in dados_saldo:
            row = tbl_sal.add_row()
            vals = [str(item['numero']), item['data_vencimento'], "R$ " + item['valor'], str(item['forma_pagamento'])]
            for i, v in enumerate(vals):
                row.cells[i].text = v
                for p in row.cells[i].paragraphs: 
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    if p.runs: aplicar_estilo(p.runs[0])
                    else: 
                        r = p.add_run(v)
                        aplicar_estilo(r)
    except Exception as e:
        print(f"Erro Tabela Saldo: {e}")

    doc.save(docx_path)

# --- GERAÇÃO PRINCIPAL ---
def gerar_contrato_pdf(aluno, turma, curso, contrato, datas_info):
    if not os.path.exists(TEMPLATE_PATH):
        st.error(f"❌ Template não encontrado: {TEMPLATE_PATH}")
        return None

    # Cálculos e Listas
    lista_entrada = datas_info.get('detalhes_entrada', [])
    lista_saldo = []
    qtd_s = int(contrato['saldo_qtd_parcelas'])
    if qtd_s > 0:
        val_s = parse_float(contrato['saldo_valor']) / qtd_s
        try: ini_s = datetime.strptime(datas_info.get('inicio_saldo'), "%Y-%m-%d")
        except: ini_s = datetime.now()
        from dateutil.relativedelta import relativedelta
        for i in range(qtd_s):
            dt = ini_s + relativedelta(months=i)
            lista_saldo.append({
                "numero": f"{i+1}/{qtd_s}",
                "data_vencimento": dt.strftime("%d/%m/%Y"),
                "valor": format_moeda(val_s),
                "forma_pagamento": contrato['saldo_forma_pagamento']
            })

    v_bruto = parse_float(contrato['valor_curso'])
    v_desc = parse_float(contrato['valor_desconto'])
    v_final = parse_float(contrato['valor_final'])
    v_material = v_bruto * 0.30 

    hoje = datetime.now()
    meses = ['Janeiro','Fevereiro','Março','Abril','Maio','Junho','Julho','Agosto','Setembro','Outubro','Novembro','Dezembro']
    
    # CONTEXTO COMPLETO (Para preencher o Word editado)
    context = {
        # Pessoais
        "nome": aluno['nome_completo'].upper(), "cpf": aluno['cpf'],
        "estado_civil": aluno.get('estado_civil', ''), "email": aluno['email'],
        "área_formação": aluno.get('area_formacao', ''), "data_nascimento": format_data(aluno.get('data_nascimento')),
        "nacionalidade": aluno.get('nacionalidade', ''), "crm": aluno.get('crm', ''),
        "telefone": aluno.get('telefone', ''), "logradouro": aluno.get('logradouro',''),
        "numero": aluno.get('numero',''), "bairro": aluno.get('bairro',''),
        "complemento": aluno.get('complemento',''), "cidade": aluno.get('cidade',''),
        "uf": aluno.get('uf',''), "cep": aluno.get('cep',''),
        
        # Produto (AGORA USADO NO WORD)
        "pos_graduacao": curso['nome'], 
        "formato_curso": contrato.get('formato_curso', 'Digital'),
        "turma": turma['codigo_turma'],
        "atendimento": "SIM" if contrato['atendimento_paciente'] else "NÃO",
        "bolsista": "SIM" if contrato['bolsista'] else "NÃO",
        "valor_curso": format_moeda(v_bruto),
        "valor_desconto": format_moeda(v_desc),
        "pencentual_desconto": f"{contrato['percentual_desconto']}%",
        "valor_final": format_moeda(v_final),
        
        # Outros
        "valor_material": format_moeda(v_material),
        "dia": hoje.day, "mês": meses[hoje.month - 1], "ano": hoje.year
    }

    try:
        # 1. Preenche TODAS as variáveis {{ }} (Incluindo a tabela de produto)
        doc = DocxTemplate(TEMPLATE_PATH)
        doc.render(context)
        
        filename = f"Contrato_{aluno['cpf']}_{turma['codigo_turma']}"
        docx_path = f"/tmp/{filename}.docx"
        pdf_path = f"/tmp/{filename}.pdf"
        doc.save(docx_path)
        
        # 2. Injeta APENAS as linhas das tabelas de pagamento (Entrada/Saldo)
        preencher_tabelas_pagamento(docx_path, lista_entrada, lista_saldo)
        
        # 3. Conversão e Upload
        cmd = ["soffice", "--headless", "--convert-to", "pdf", "--outdir", "/tmp", docx_path]
        subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        
        final_local = pdf_path if os.path.exists(pdf_path) else docx_path
        mime = "application/pdf" if os.path.exists(pdf_path) else "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = ".pdf" if os.path.exists(pdf_path) else ".docx"
        
        path_cloud = f"{hoje.year}/{hoje.month}/{filename}{ext}"
        with open(final_local, "rb") as f:
            supabase.storage.from_("contratos").upload(path_cloud, f, {"content-type": mime, "upsert": "true"})
            
        return final_local, path_cloud

    except Exception as e:
        st.error(f"Erro: {e}")
        return None

# --- MANTER FUNÇÕES DE EMAIL/CARIMBO ---
def aplicar_carimbo_digital(path, meta):
    # (Manter código igual ao anterior)
    if not path.endswith(".pdf"): return path
    try:
        d = supabase.storage.from_("contratos").download(path)
        r = PdfReader(io.BytesIO(d)); w = PdfWriter()
        p = io.BytesIO(); c = canvas.Canvas(p, pagesize=letter)
        c.setFont("Helvetica",6); c.setFillColorRGB(0.5,0.5,0.5,0.5)
        txt = f"ACEITE DIGITAL | {meta['data_hora']} | {meta['nome']} | CPF:{meta['cpf']} | IP:{meta['ip']} | Hash:{meta['hash'][:10]}..."
        c.drawString(20,20,txt); c.save(); p.seek(0)
        wm = PdfReader(p).pages[0]
        for pg in r.pages: pg.merge_page(wm); w.add_page(pg)
        out = io.BytesIO(); w.write(out); out.seek(0)
        np = path.replace(".pdf","_assinado.pdf")
        supabase.storage.from_("contratos").upload(np, out, {"content-type":"application/pdf","upsert":"true"})
        return np
    except: return path

def enviar_email(dest, nome, link):
    # (Manter código igual ao anterior)
    try:
        msg = MIMEMultipart()
        msg['From']=st.secrets["GMAIL_EMAIL"]; msg['To']=dest; msg['Subject']="Assinatura NexusMed"
        html = f"""<div style="font-family:Arial;padding:20px;border:1px solid #ddd;border-radius:5px;"><h2 style="color:#004B8D;">Olá, {nome}</h2><p>Seu contrato está disponível.</p><br><a href="{link}" style="background:#004B8D;color:#fff;padding:12px 24px;text-decoration:none;border-radius:4px;">ASSINAR CONTRATO</a></div>"""
        msg.attach(MIMEText(html, 'html'))
        s=smtplib.SMTP("smtp.gmail.com",587); s.starttls(); s.login(st.secrets["GMAIL_EMAIL"],st.secrets["GMAIL_PASSWORD"])
        s.sendmail(msg['From'],dest,msg.as_string()); s.quit(); return True
    except: return False
