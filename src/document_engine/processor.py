import io
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

class ContractProcessor:
    """
    Motor de processamento de documentos.
    Responsável por transformar o template .docx em um contrato preenchido.
   
    """

    def __init__(self, template_path: str):
        self.template_path = template_path

    def generate_docx(self, context: dict, entry_rows: list, installment_rows: list) -> io.BytesIO:
        """
        Gera o documento preenchido e retorna um buffer de memória (BytesIO).
        context: Dicionário com as variáveis {{ nome }}, {{ cpf }}, etc.
        entry_rows: Lista para a tabela de ENTRADA (Seção 03).
        installment_rows: Lista para a tabela de SALDO (Seção 04).
       
        """
        # 1. Primeiro usamos o DocxTemplate para variáveis simples
        doc_tpl = DocxTemplate(self.template_path)
        doc_tpl.render(context)
        
        # Salva em um buffer temporário para a segunda fase
        temp_buffer = io.BytesIO()
        doc_tpl.save(temp_buffer)
        temp_buffer.seek(0)

        # 2. Usamos python-docx para manipular as tabelas de forma dinâmica
        return self._inject_payment_tables(temp_buffer, entry_rows, installment_rows)

    def _inject_payment_tables(self, doc_stream: io.BytesIO, entry_data: list, installment_data: list) -> io.BytesIO:
        """
        Localiza as tabelas de Entrada e Saldo e adiciona as linhas dinamicamente.
       
        """
        doc = Document(doc_stream)
        
        # Identifica tabelas que possuem "Vencimento" no cabeçalho
        payment_tables = []
        for table in doc.tables:
            try:
                # Verifica se a segunda célula contém o termo chave (Vencimento)
                if "Vencimento" in table.rows[0].cells[1].text:
                    payment_tables.append(table)
            except:
                continue
        
        # Preenche a Tabela de Entrada (Primeira encontrada - Seção 03)
        if len(payment_tables) > 0 and entry_data:
            self._fill_table_rows(payment_tables[0], entry_data)

        # Preenche a Tabela de Saldo (Segunda encontrada - Seção 04)
        if len(payment_tables) > 1 and installment_data:
            self._fill_table_rows(payment_tables[1], installment_data)

        final_buffer = io.BytesIO()
        doc.save(final_buffer)
        final_buffer.seek(0)
        return final_buffer

    def _fill_table_rows(self, table, data_rows: list):
        """
        Limpa linhas extras e adiciona dados com formatação Arial 10 e Centralizado.
       
        """
        # 1. Limpeza: Mantém apenas o cabeçalho (remove linhas residuais do Word)
        while len(table.rows) > 1:
            table._element.remove(table.rows[-1]._element)

        # 2. Injeção de Dados
        for item in data_rows:
            row = table.add_row()
            
            # Mapeia os dados para as 4 colunas padrão do contrato
            # Espera-se que cada item tenha: parcela, vencimento, valor, forma
            valores = [
                str(item.get('parcela', '')),
                str(item.get('vencimento', '')),
                str(item.get('valor', '')),
                str(item.get('forma', ''))
            ]

            for i, texto in enumerate(valores):
                cell = row.cells[i]
                paragraph = cell.paragraphs[0]
                paragraph.clear() # Remove qualquer formatação herdada
                
                run = paragraph.add_run(texto)
                
                # --- APLICAÇÃO DE ESTILO RIGOROSO ---
                run.font.name = 'Arial'
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
